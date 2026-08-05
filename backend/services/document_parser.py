"""
IntelliReal - Document Parser
Extracts text content from multiple file formats: PDF, DOCX, XLSX, CSV, TXT, HTML.
Returns structured text with page/section metadata for downstream chunking.
"""

import io
import logging
from typing import List, Dict, Optional
from pathlib import Path

import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


class ParsedDocument:
    """Represents a parsed document with extracted text and metadata."""

    def __init__(self, filename: str, file_type: str):
        self.filename = filename
        self.file_type = file_type
        self.pages: List[Dict] = []  # [{"page": 1, "text": "..."}, ...]
        self.tables: List[Dict] = []  # [{"page": 1, "data": [[...]], "headers": [...]}]
        self.metadata: Dict = {}

    @property
    def full_text(self) -> str:
        """Combine all pages into a single string."""
        return "\n\n".join(page["text"] for page in self.pages if page["text"].strip())

    @property
    def total_pages(self) -> int:
        return len(self.pages)

    def to_dict(self) -> dict:
        return {
            "filename": self.filename,
            "file_type": self.file_type,
            "total_pages": self.total_pages,
            "tables_found": len(self.tables),
            "metadata": self.metadata,
        }


class DocumentParser:
    """
    Multi-format document parser.
    Supports: PDF, DOCX, XLSX, CSV, TXT, HTML
    """

    SUPPORTED_TYPES = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".xlsx": "xlsx",
        ".csv": "csv",
        ".txt": "txt",
        ".html": "html",
        ".htm": "html",
    }

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """
        Parse a document from raw bytes.
        
        Args:
            file_bytes: Raw file content
            filename: Original filename (used to determine type)
            
        Returns:
            ParsedDocument with extracted text and metadata
        """
        ext = Path(filename).suffix.lower()
        file_type = self.SUPPORTED_TYPES.get(ext)

        if not file_type:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(self.SUPPORTED_TYPES.keys())}"
            )

        logger.info(f"Parsing {filename} as {file_type}")

        parser_map = {
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "xlsx": self._parse_xlsx,
            "csv": self._parse_csv,
            "txt": self._parse_txt,
            "html": self._parse_html,
        }

        doc = ParsedDocument(filename=filename, file_type=file_type)
        parser_map[file_type](file_bytes, doc)

        logger.info(
            f"Parsed {filename}: {doc.total_pages} pages, "
            f"{len(doc.tables)} tables, "
            f"{len(doc.full_text)} chars"
        )

        return doc

    def _parse_pdf(self, file_bytes: bytes, doc: ParsedDocument) -> None:
        """Extract text from PDF, page by page."""
        reader = PdfReader(io.BytesIO(file_bytes))

        doc.metadata = {
            "num_pages": len(reader.pages),
            "info": {k: str(v) for k, v in (reader.metadata or {}).items()},
        }

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            doc.pages.append({"page": i + 1, "text": text})

            # Attempt basic table detection from text patterns
            if "|" in text or "\t" in text:
                self._extract_text_tables(text, i + 1, doc)

    def _parse_docx(self, file_bytes: bytes, doc: ParsedDocument) -> None:
        """Extract text from Word documents."""
        docx_doc = DocxDocument(io.BytesIO(file_bytes))

        # Extract paragraphs
        full_text = []
        for para in docx_doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        doc.pages.append({"page": 1, "text": "\n".join(full_text)})

        # Extract tables
        for i, table in enumerate(docx_doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if rows:
                doc.tables.append({
                    "page": 1,
                    "table_index": i,
                    "headers": rows[0] if rows else [],
                    "data": rows[1:] if len(rows) > 1 else rows,
                })

        doc.metadata = {"num_tables": len(docx_doc.tables)}

    def _parse_xlsx(self, file_bytes: bytes, doc: ParsedDocument) -> None:
        """Extract data from Excel spreadsheets."""
        xls = pd.ExcelFile(io.BytesIO(file_bytes))

        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)

            # Convert sheet to text representation
            text_lines = [f"Sheet: {sheet_name}", "=" * 40]
            text_lines.append(df.to_string(index=False))

            doc.pages.append({
                "page": xls.sheet_names.index(sheet_name) + 1,
                "text": "\n".join(text_lines),
            })

            # Store as structured table
            doc.tables.append({
                "page": xls.sheet_names.index(sheet_name) + 1,
                "sheet": sheet_name,
                "headers": list(df.columns),
                "data": df.values.tolist(),
            })

        doc.metadata = {"sheets": xls.sheet_names}

    def _parse_csv(self, file_bytes: bytes, doc: ParsedDocument) -> None:
        """Extract data from CSV files."""
        df = pd.read_csv(io.BytesIO(file_bytes))

        text = df.to_string(index=False)
        doc.pages.append({"page": 1, "text": text})

        doc.tables.append({
            "page": 1,
            "headers": list(df.columns),
            "data": df.values.tolist(),
        })

        doc.metadata = {"rows": len(df), "columns": list(df.columns)}

    def _parse_txt(self, file_bytes: bytes, doc: ParsedDocument) -> None:
        """Extract text from plain text files."""
        text = file_bytes.decode("utf-8", errors="replace")
        doc.pages.append({"page": 1, "text": text})

    def _parse_html(self, file_bytes: bytes, doc: ParsedDocument) -> None:
        """Extract text and tables from HTML documents."""
        soup = BeautifulSoup(file_bytes, "lxml")

        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()

        # Extract main text
        text = soup.get_text(separator="\n", strip=True)
        doc.pages.append({"page": 1, "text": text})

        # Extract HTML tables
        for i, table in enumerate(soup.find_all("table")):
            rows = []
            for tr in table.find_all("tr"):
                cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                rows.append(cells)

            if rows:
                doc.tables.append({
                    "page": 1,
                    "table_index": i,
                    "headers": rows[0] if rows else [],
                    "data": rows[1:] if len(rows) > 1 else rows,
                })

        # Extract metadata
        title = soup.find("title")
        doc.metadata = {
            "title": title.get_text(strip=True) if title else None,
            "num_tables": len(soup.find_all("table")),
        }

    def _extract_text_tables(
        self, text: str, page_num: int, doc: ParsedDocument
    ) -> None:
        """Best-effort table extraction from text patterns (pipe-delimited, tab-delimited)."""
        lines = text.split("\n")
        table_rows = []
        in_table = False

        for line in lines:
            if "|" in line and line.count("|") >= 2:
                cells = [c.strip() for c in line.split("|") if c.strip()]
                if cells:
                    table_rows.append(cells)
                    in_table = True
            elif in_table and table_rows:
                # End of table block
                doc.tables.append({
                    "page": page_num,
                    "headers": table_rows[0],
                    "data": table_rows[1:],
                })
                table_rows = []
                in_table = False

        # Flush remaining
        if table_rows:
            doc.tables.append({
                "page": page_num,
                "headers": table_rows[0],
                "data": table_rows[1:],
            })
